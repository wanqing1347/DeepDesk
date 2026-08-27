from io import BytesIO

import pytest
from docx import Document as DocxDocument
from pypdf import PdfWriter

from app.files.parser import FileParser


def test_txt_parser_preserves_unicode_and_applies_display_truncation_only() -> None:
    parser = FileParser(max_text_chars=5)

    result = parser.parse(file_name="notes.txt", content="你好abcdef".encode())

    assert result.full_text == "你好abcdef"
    assert result.truncated_text == "你好abc\n\n... (内容已截断，文件过长)"


def test_docx_parser_extracts_non_blank_paragraphs() -> None:
    buffer = BytesIO()
    document = DocxDocument()
    document.add_paragraph("第一段")
    document.add_paragraph("   ")
    document.add_paragraph("第二段")
    document.save(buffer)

    result = FileParser().parse(file_name="document.docx", content=buffer.getvalue())

    assert result.full_text == "第一段\n第二段"
    assert result.truncated_text == result.full_text


def test_pdf_parser_accepts_valid_pdf_and_returns_extracted_text() -> None:
    buffer = BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    writer.write(buffer)

    result = FileParser().parse(file_name="blank.pdf", content=buffer.getvalue())

    assert result.full_text == ""
    assert result.truncated_text == ""


def test_legacy_doc_is_explicitly_unsupported() -> None:
    with pytest.raises(ValueError, match=r"暂不支持 \.doc 格式"):
        FileParser().parse(file_name="legacy.doc", content=b"legacy")


def test_unknown_extension_is_rejected() -> None:
    with pytest.raises(ValueError, match="不支持的文件类型: md"):
        FileParser().parse(file_name="readme.md", content=b"markdown")
